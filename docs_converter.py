"""
   CodeCraft PMS Backend Project

   파일명   : docs_converter.py                                                          
   생성자   : 김창환                                                          
                                                                               
   생성일   : 2024/11/26                                                      
   업데이트 : 2024/12/03                                                       
                                                                               
   설명     : DB로부터 정보를 받아와 산출물을 문서화 해주는 기능 정의
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime, date
import pymysql, os, sys, traceback
import re  # 정규식 사용

sys.path.append(os.path.abspath('/data/Database Project'))
import output_DB

router = APIRouter()

class ConverterPayload(BaseModel):
    doc_type: int
    doc_s_no: int

def replace_placeholder_in_cell(cell, placeholder, replacement):
    """
    셀 내부의 자리표시자를 치환합니다.
    모든 Run을 삭제 후 하나의 Run으로 대체합니다.
    """
    for paragraph in cell.paragraphs:
        # 모든 Run 텍스트 병합
        full_text = "".join(run.text for run in paragraph.runs)

        # 치환 대상이 있을 경우
        if placeholder in full_text:
            # 치환된 텍스트 생성
            updated_text = full_text.replace(placeholder, replacement)

            # 기존 Run 삭제
            for run in paragraph.runs:
                run.clear()

            # 새로운 Run으로 대체
            paragraph.clear()  # 전체 문단 초기화
            paragraph.add_run(updated_text)

@router.post("/docs/convert")
async def docs_convert(payload: ConverterPayload):
    try:
        if payload.doc_type == 1:  # 회의록
            meeting_data = output_DB.fetch_one_meeting_minutes(payload.doc_s_no)
            if not meeting_data:
                raise HTTPException(status_code=404, detail="Meeting data not found")

            try:
                doc = Document("/data/Docs_Template/회의록.docx")
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Template loading error: {e}")

            # 자리표시자 치환
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_placeholder_in_cell(cell, "{f1}", meeting_data.get("doc_m_title", ""))
                        replace_placeholder_in_cell(
                            cell, "{f2}",
                            meeting_data.get("doc_m_date", "").strftime("%Y-%m-%d")
                            if isinstance(meeting_data.get("doc_m_date"), date)
                            else str(meeting_data.get("doc_m_date", ""))
                        )
                        replace_placeholder_in_cell(cell, "{f3}", meeting_data.get("doc_m_manager", ""))
                        replace_placeholder_in_cell(cell, "{f4}", meeting_data.get("doc_m_loc", ""))
                        replace_placeholder_in_cell(cell, "{f5}", str(len(meeting_data.get("doc_m_member", "").split(";"))))
                        replace_placeholder_in_cell(cell, "{f6}", meeting_data.get("doc_m_content", ""))
                        replace_placeholder_in_cell(cell, "{f7}", meeting_data.get("doc_m_result", ""))
                        replace_placeholder_in_cell(cell, "{f8}", "김창환")
            # # 참석자 데이터 처리
            # members = meeting_data.get("doc_m_member", "").split(";")
            # for i, member in enumerate(members):
            #     if i >= 4:  # 템플릿에 최대 4명까지만 지원
            #         break
            #     name_match = re.match(r"([가-힣]+)(\d+)", member)
            #     if name_match:
            #         name, student_id = name_match.groups()
            #         replace_placeholder_in_cell(doc.tables[0].cell(i + 1, 0), f"{{f{i * 2 + 8}}}", name)
            #         replace_placeholder_in_cell(doc.tables[0].cell(i + 1, 1), f"{{f{i * 2 + 9}}}", student_id)

            # # 나머지 자리표시자 초기화
            # for i in range(8, 16):
            #     placeholder = f"{{f{i}}}"
            #     for table in doc.tables:
            #         for row in table.rows:
            #             for cell in row.cells:
            #                 replace_placeholder_in_cell(cell, placeholder, "")

            # 문서 저장
            output_path = f"/data/Backend Project/temp/회의록_{payload.doc_s_no}.docx"
            doc.save(output_path)

            return {"RESULT_CODE": 200, "RESULT_MSG": "Done!", "FILE_PATH": output_path}

    except Exception:
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail="Unexpected error occurred")